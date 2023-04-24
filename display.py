import matplotlib.pyplot as plt
import seaborn as sns
import matplotlib as mpl
import numpy as np
import pandas as pd
import xlsxwriter
import docx

#For bidirectional, horizontal data chart
pd.options.mode.chained_assignment = None  # default='warn'
facecolor = '#eaeaf2'
color_red = '#fd625e'
color_blue = '#3333ff'
font_color = '#525252'
hfont = {'fontname':'Calibri'}
display_num = 50
listing_num = 200
cancer_dict = {
    'breast' : "vú",
    'lung' : "phổi",
    'thyroid' : "giáp",
    'large_intestine' : "đại tràng",
    'hepatocellular_carcinoma' : "gan"
}
np.seterr(all="ignore")

def _load_data(data, index_len): 
    data1 = pd.DataFrame(
        {'list' : data['list'],
         'ratio' : data['ratio']
        }).head(len(data['list']))
    data1 = data1.sort_values(['ratio'], ascending=False).reset_index(drop=True)
    data = pd.DataFrame(
        {'list' : data1['list'],
         'ratio' : data1['ratio']
        }).head(index_len)
    return data

def _load_ticks(data, index_len):
    data['ratio'] = [x*100 for x in data['ratio'][0:index_len]]
    for x in range(0,index_len):
        data['list'][x] += f" ({round(data['ratio'][x], 2)}%)"
    return data

def write_to_doc(asia_gene, world_gene, cancer_name):
    mydoc = docx.Document()
    mydoc.add_paragraph(f"Potential panel genes of {cancer_name} cancer are (n={listing_num}):")
    para = "Asia:\n"
    data = pd.DataFrame(
        {'list' : asia_gene['list'],
         'ratio' : asia_gene['ratio']
        }).head(len(asia_gene['list']))
    data = data.sort_values('ratio', ascending=False).reset_index(drop=True)
    data.to_csv(f"output/{cancer_name}_cancer/drug/panel_gene_asia.csv")

    for k in range(0, min(len(data), listing_num)):
        para += f"{k}, {data['list'][k]} ({round(data['ratio'][k]*100,4)}%)\n"
    para += "\n\n"
    para += "The world:\n"
    data = pd.DataFrame(
        {'list' : world_gene['list'],
         'ratio' : world_gene['ratio']
        }).head(len(world_gene['list']))
    data = data.sort_values('ratio', ascending=False).reset_index(drop=True)
    data.to_csv(f"output/{cancer_name}_cancer/drug/panel_gene_world.csv")
    for k in range(0, min(len(data), listing_num)):
        para += f"{k}, {data['list'][k]} ({round(data['ratio'][k]*100,4)}%)\n"
    mydoc.add_paragraph(para)
    mydoc.save(f"output/panels/{cancer_name}_cancer_gene_list.docx")



def _chart_create(ratio, ratio_census, cancer_name, flag, title0, title1):
    sns.set_theme()
    #load data
    index_len = min(display_num, len(ratio['list']), len(ratio_census['list'])) 
    data_asian = _load_data(ratio, index_len)
    data_census = _load_data(ratio_census, index_len)
    data_asian = _load_ticks(data_asian, index_len)
    data_census = _load_ticks(data_census, index_len)
    #set up data
    tick_limit = min(10, round(max(data_asian['ratio'].iloc[0], data_census['ratio'].iloc[0]) / 10 + 1))
    tick_value = [1]*tick_limit
    for x in range (1, tick_limit+1):    tick_value[x-1] = x*10
    data_asian.set_index('list', inplace=True)
    data_census.set_index('list', inplace=True)
    indexa = data_asian.index
    indexw = data_census.index
    fig, axes = plt.subplots(figsize=(17,8), facecolor=facecolor, ncols=2, squeeze=True)
    #set up individual bars
    axes[0].barh(indexa, data_asian['ratio'], alpha = 0.8, align='center', color=color_red, zorder=2)
    axes[0].set_title(title0, fontsize=14, pad=8, color=color_red, **hfont)
    axes[1].barh(indexa, data_census['ratio'], alpha = 0.8, align='center', color=color_blue, zorder=2)
    axes[1].set_title(title1, fontsize=14, pad=8, color=color_blue, **hfont)
    axes[0].invert_xaxis()
    axes[0].invert_yaxis()
    axes[1].invert_yaxis()
    #set up ticks and tick labels
    axes[0].set(yticks=data_asian.index, yticklabels=indexa)
    axes[0].yaxis.tick_left()
    axes[0].tick_params(axis='y', colors='grey') 
 

    axes[1].set_yticklabels(indexw)
    axes[1].yaxis.set_ticks_position("right")
    axes[1].tick_params(axis='y', colors='grey') 
    
    axes[0].set_xticks(tick_value)
    axes[0].set_xticklabels(tick_value)

    axes[1].set_xticks(tick_value)
    axes[1].set_xticklabels(tick_value)

    for label in (axes[0].get_xticklabels() + axes[0].get_yticklabels()):
        label.set(fontsize=8, color=font_color, **hfont)
    for label in (axes[1].get_xticklabels() + axes[1].get_yticklabels()):
        label.set(fontsize=8, color=font_color, **hfont)
    plt.subplots_adjust(wspace=0, top=0.85, bottom=0.1, left=0.18, right=0.95)
    if flag == 'g':
        filename = f"output/{cancer_name}_cancer/charts/{cancer_name}_cancer"
        plt.xlabel(f'Top {index_len} gen có xác nhận đột biến của ung thư {cancer_dict[cancer_name]} theo tần suất (%) (COSMIC)',        
    x = -0.04,fontsize = 12, labelpad=6)
    else:
        filename = f"output/{cancer_name}_cancer/charts/(variant){cancer_name}_gene"
        plt.xlabel(f'Top {index_len} AA variants among samples of {cancer_name.replace("_", " ")} cancer with confirmed mutations (%) (extracted from CosmicMutantExportCensus.tsv)', x = -0.04,fontsize = 12, labelpad=6)    
    #plt.title(label=f'Ratios of top 50 genes among mutated samples of {cancer_name} cancer', y = -0.15, loc="below")
    plt.legend()
    #plt.rcParams["figure.autolayout"] = True
    plt.rcParams["figure.figsize"] = [20, 10]
    #plt.tight_layout()
    plt.savefig(filename+'.png', dpi = 250,bbox_inches='tight', facecolor=facecolor)
    #plt.show()

